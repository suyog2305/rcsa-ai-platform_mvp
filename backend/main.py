"""
RCSA Intelligence Platform — Backend
═══════════════════════════════════════════════════════════════════════
Real RAG pipeline: Word docs → ChromaDB → GPT-4o → Sourced responses

Architecture:
  1. On startup → load all .docx files from /control_docs
  2. Chunk + embed via OpenAI text-embedding-3-small → store in ChromaDB
  3. Each chat message → semantic search → retrieve top 4 chunks
  4. Inject chunks as context → GPT-4o generates response
  5. Return response + source attribution to frontend
═══════════════════════════════════════════════════════════════════════
"""
import os
import logging
from pathlib import Path
from typing import List, Optional
from datetime import datetime

from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from openai import OpenAI
import chromadb
from chromadb.utils import embedding_functions
from docx import Document as DocxDocument

# ── Configuration ───────────────────────────────────────────────────────
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
if not OPENAI_API_KEY:
    raise ValueError("OPENAI_API_KEY environment variable not set")

DOCS_PATH = Path(__file__).parent / "control_docs"
CHROMA_PATH = Path(__file__).parent / "chroma_data"
COLLECTION_NAME = "rcsa_controls"
CHAT_MODEL = "gpt-4o"
EMBEDDING_MODEL = "text-embedding-3-small"
CHUNK_SIZE = 800
CHUNK_OVERLAP = 120
TOP_K_RETRIEVAL = 4

# Control-gap recommendations ship dark. Set ENABLE_CONTROL_RECO=true in the
# deployment's environment variables to expose /control_recommendations.
ENABLE_CONTROL_RECO = os.getenv("ENABLE_CONTROL_RECO", "false").strip().lower() in ("1", "true", "yes", "on")
CONTROL_GAP_TOP_K = 10

# Voice input via Azure Speech. Needs the flag AND a key/region — an enabled flag with
# no resource behind it would show the control owner a mic button that cannot work, so
# both are required before the frontend is told the feature exists.
ENABLE_VOICE = os.getenv("ENABLE_VOICE", "false").strip().lower() in ("1", "true", "yes", "on")
AZURE_SPEECH_KEY = os.getenv("AZURE_SPEECH_KEY", "").strip()
AZURE_SPEECH_REGION = os.getenv("AZURE_SPEECH_REGION", "").strip()
AZURE_SPEECH_LANGUAGE = os.getenv("AZURE_SPEECH_LANGUAGE", "en-GB").strip()
VOICE_AVAILABLE = ENABLE_VOICE and bool(AZURE_SPEECH_KEY and AZURE_SPEECH_REGION)
# Azure issues a 10-minute token; expire ours earlier so the browser refreshes in time.
SPEECH_TOKEN_TTL_SECONDS = 540

logging.basicConfig(level=logging.INFO, format="%(asctime)s · %(levelname)s · %(message)s")
log = logging.getLogger("rcsa")

# ── Clients ─────────────────────────────────────────────────────────────
openai_client = OpenAI(api_key=OPENAI_API_KEY)
chroma_client = chromadb.PersistentClient(path=str(CHROMA_PATH))
openai_ef = embedding_functions.OpenAIEmbeddingFunction(
    api_key=OPENAI_API_KEY,
    model_name=EMBEDDING_MODEL,
)

# ── FastAPI app ─────────────────────────────────────────────────────────
app = FastAPI(
    title="RCSA Intelligence Platform API",
    description="Real RAG over banking control documents · GPT-4o + ChromaDB",
    version="1.0.0",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # In production, restrict to your Vercel domain
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ── Document loading & chunking ─────────────────────────────────────────
def load_docx(path: Path) -> str:
    """Extract all text from a .docx file."""
    doc = DocxDocument(str(path))
    parts = []
    # Paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    # Tables (control docs have many)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def chunk_text(text: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    """Sliding-window chunking with paragraph awareness."""
    if len(text) <= size:
        return [text]
    chunks = []
    start = 0
    while start < len(text):
        end = min(start + size, len(text))
        # Try to break on paragraph boundary
        if end < len(text):
            last_break = text.rfind("\n", start, end)
            if last_break > start + size // 2:
                end = last_break
        chunks.append(text[start:end].strip())
        start = end - overlap if end < len(text) else end
    return [c for c in chunks if len(c) > 50]


def ingest_documents():
    """Ingest all .docx files into ChromaDB on startup."""
    # Reset collection (idempotent on cold starts)
    try:
        chroma_client.delete_collection(COLLECTION_NAME)
        log.info("Existing collection deleted")
    except Exception:
        pass

    collection = chroma_client.create_collection(
        name=COLLECTION_NAME,
        embedding_function=openai_ef,
        metadata={"hnsw:space": "cosine"},
    )

    if not DOCS_PATH.exists():
        log.warning(f"Docs path {DOCS_PATH} does not exist — skipping ingestion")
        return collection

    doc_files = sorted(DOCS_PATH.glob("*.docx"))
    log.info(f"Found {len(doc_files)} documents to ingest")

    all_chunks = []
    all_ids = []
    all_metadata = []

    for doc_path in doc_files:
        try:
            text = load_docx(doc_path)
            chunks = chunk_text(text)
            doc_name = doc_path.stem
            # Parse control ID from filename (e.g. RB-CTRL-02_Control_Narrative)
            parts = doc_name.split("_", 1)
            control_id = parts[0] if parts else doc_name
            doc_type = parts[1].replace("_", " ") if len(parts) > 1 else "Document"

            for i, chunk in enumerate(chunks):
                all_chunks.append(chunk)
                all_ids.append(f"{doc_name}_chunk_{i}")
                all_metadata.append({
                    "source": doc_name,
                    "control_id": control_id,
                    "doc_type": doc_type,
                    "chunk_index": i,
                })

            log.info(f"  ✓ {doc_name}: {len(chunks)} chunks")
        except Exception as e:
            log.error(f"  ✗ Failed to ingest {doc_path.name}: {e}")

    if all_chunks:
        # Batch insert (Chroma handles embedding via the embedding_function)
        batch_size = 100
        for i in range(0, len(all_chunks), batch_size):
            collection.add(
                documents=all_chunks[i:i+batch_size],
                ids=all_ids[i:i+batch_size],
                metadatas=all_metadata[i:i+batch_size],
            )
        log.info(f"✓ Ingested {len(all_chunks)} total chunks into vector store")

    return collection


# ── Initialise collection on startup ────────────────────────────────────
collection = None

@app.on_event("startup")
async def startup_event():
    global collection
    log.info("═" * 60)
    log.info("RCSA Intelligence Platform — Starting up")
    log.info("═" * 60)
    collection = ingest_documents()
    log.info("✓ Ready to serve requests")


# ── Pydantic models ─────────────────────────────────────────────────────
class ChatMessage(BaseModel):
    role: str  # "user" or "assistant"
    content: str


class ChatRequest(BaseModel):
    messages: List[ChatMessage]
    business_unit: Optional[str] = "Retail Banking"


class RetrievedChunk(BaseModel):
    source: str
    control_id: str
    doc_type: str
    snippet: str
    similarity: float


class ChatResponse(BaseModel):
    response: str
    retrieved_chunks: List[RetrievedChunk]
    model: str
    timestamp: str


class RCSAOutputRequest(BaseModel):
    messages: List[ChatMessage]
    business_unit: Optional[str] = "Retail Banking"


class RCSAOutputResponse(BaseModel):
    business_unit: str
    inherent_risk_score: int
    inherent_risk_label: str
    residual_risk_score: int
    residual_risk_label: str
    priority_action_count: int
    heat_map_position: dict  # {impact: 1-5, likelihood: 1-5}
    detective_controls: List[dict]
    three_lines_defence: dict
    priority_actions: List[dict]
    ai_rationale: str
    evidence_trace: List[str]
    generated_at: str


class ControlGapRequest(BaseModel):
    messages: List[ChatMessage]
    business_unit: Optional[str] = "Retail Banking"


class ControlGapResponse(BaseModel):
    business_unit: str
    unmitigated_risk_count: int
    inadequate_control_count: int
    unmitigated_risks: List[dict]
    inadequate_controls: List[dict]
    coverage_summary: str
    evidence_trace: List[str]
    generated_at: str


class DashboardResponse(BaseModel):
    business_units_in_scope: int
    high_risk_bus: int
    open_action_items: int
    controls_rated_effective_pct: int
    business_units: List[dict]
    top_risk_categories: List[dict]
    cycle_time_comparison: dict
    generated_at: str


# ── System prompt ───────────────────────────────────────────────────────
SYSTEM_PROMPT = """You are an expert RCSA (Risk and Control Self-Assessment) AI Agent for a large global retail bank. You conduct structured risk interviews with control owners as part of the quarterly RCSA cycle.

You have access to the bank's full library of control documentation via semantic retrieval — control narratives, performance evaluations, design testing evidence, owner interview transcripts. Each user message triggers a search across this knowledge base and the relevant chunks are injected below.

INTERVIEW GUIDELINES:
1. Conduct a professional, structured RCSA interview — adaptive, not scripted
2. Ask one focused question at a time, with intelligent follow-up
3. When you identify risks, flag them explicitly with severity (HIGH/MEDIUM/LOW) and the specific regulatory citation
4. Reference specific incidents, control IDs, regulations, and dates from the retrieved knowledge
5. After 6-8 exchanges, produce a structured assessment summary with inherent risk, residual risk, and 3 priority actions
6. Keep responses concise — max 3-4 sentences per turn unless producing a summary
7. Be professional and respectful — you are assisting the risk owner, not interrogating

OUTPUT FORMAT:
- For risk flags, use: **Risk flagged — [HIGH/MEDIUM/LOW]** with regulatory citation
- For final summaries, structure as: Inherent Risk Score (x/25), Residual Risk Score (x/25), Top 3 Priority Actions
"""


# ── Routes ──────────────────────────────────────────────────────────────
@app.get("/")
async def root():
    return {
        "service": "RCSA Intelligence Platform",
        "status": "operational",
        "model": CHAT_MODEL,
        "documents_indexed": collection.count() if collection else 0,
        "features": {
            "control_recommendations": ENABLE_CONTROL_RECO,
            "voice_input": VOICE_AVAILABLE,
        },
    }


@app.get("/health")
async def health():
    return {"status": "healthy", "timestamp": datetime.utcnow().isoformat()}


@app.get("/knowledge_base")
async def list_kb():
    """List all documents currently in the knowledge base."""
    if not collection:
        raise HTTPException(503, "Knowledge base not initialised")

    # Get unique sources
    all_items = collection.get(include=["metadatas"])
    sources = {}
    for meta in all_items["metadatas"]:
        src = meta["source"]
        if src not in sources:
            sources[src] = {
                "source": src,
                "control_id": meta["control_id"],
                "doc_type": meta["doc_type"],
                "chunk_count": 0,
            }
        sources[src]["chunk_count"] += 1

    return {
        "total_documents": len(sources),
        "total_chunks": len(all_items["metadatas"]),
        "documents": sorted(sources.values(), key=lambda x: x["source"]),
    }


@app.post("/chat", response_model=ChatResponse)
async def chat(request: ChatRequest):
    """Main RAG endpoint — retrieve relevant chunks + generate response."""
    if not collection:
        raise HTTPException(503, "Knowledge base not initialised")

    # Get the latest user message
    user_messages = [m for m in request.messages if m.role == "user"]
    if not user_messages:
        raise HTTPException(400, "No user message provided")
    latest_query = user_messages[-1].content

    # ── 1. Retrieve relevant chunks ──
    try:
        results = collection.query(
            query_texts=[latest_query],
            n_results=TOP_K_RETRIEVAL,
            include=["documents", "metadatas", "distances"],
        )
    except Exception as e:
        log.error(f"Retrieval failed: {e}")
        raise HTTPException(500, f"Retrieval error: {str(e)}")

    retrieved_chunks = []
    context_block = ""
    if results["documents"] and results["documents"][0]:
        for doc, meta, dist in zip(
            results["documents"][0],
            results["metadatas"][0],
            results["distances"][0],
        ):
            similarity = max(0.0, 1.0 - dist)  # cosine distance → similarity
            retrieved_chunks.append(RetrievedChunk(
                source=meta["source"],
                control_id=meta["control_id"],
                doc_type=meta["doc_type"],
                snippet=doc[:300] + ("..." if len(doc) > 300 else ""),
                similarity=round(similarity, 3),
            ))
            context_block += f"\n--- From {meta['source']} ({meta['doc_type']}) ---\n{doc}\n"

    # ── 2. Build LLM messages ──
    llm_messages = [
        {"role": "system", "content": SYSTEM_PROMPT},
        {"role": "system", "content": f"RETRIEVED KNOWLEDGE BASE CONTEXT:\n{context_block}\n\n--- END RETRIEVED CONTEXT ---"},
    ]
    # Include conversation history
    for msg in request.messages:
        llm_messages.append({"role": msg.role, "content": msg.content})

    # ── 3. Generate response ──
    try:
        completion = openai_client.chat.completions.create(
            model=CHAT_MODEL,
            messages=llm_messages,
            max_tokens=600,
            temperature=0.7,
        )
        response_text = completion.choices[0].message.content
    except Exception as e:
        log.error(f"LLM call failed: {e}")
        raise HTTPException(500, f"LLM error: {str(e)}")

    log.info(f"Chat · {len(retrieved_chunks)} chunks retrieved · query: {latest_query[:80]}...")

    return ChatResponse(
        response=response_text,
        retrieved_chunks=retrieved_chunks,
        model=CHAT_MODEL,
        timestamp=datetime.utcnow().isoformat(),
    )


# ── RCSA OUTPUT GENERATION ──────────────────────────────────────────────
RCSA_OUTPUT_PROMPT = """You are generating a structured RCSA (Risk and Control Self-Assessment) output for the Risk Committee based on the interview conversation provided.

Analyse the conversation and produce a JSON object with this EXACT structure (no markdown, no extra text, just valid JSON):

{
  "inherent_risk_score": <integer 1-25>,
  "inherent_risk_label": "<HIGH|MEDIUM-HIGH|MEDIUM|MEDIUM-LOW|LOW>",
  "residual_risk_score": <integer 1-25>,
  "residual_risk_label": "<HIGH|MEDIUM-HIGH|MEDIUM|MEDIUM-LOW|LOW>",
  "priority_action_count": <integer>,
  "heat_map_position": {"impact": <1-5>, "likelihood": <1-5>},
  "detective_controls": [
    {"name": "<control name>", "type": "<Detective|Preventive>", "description": "<short desc>", "rating": "<Strong|Adequate|Weak|Ineffective>", "last_tested": "<date and source>", "notes": "<optional notes>"}
  ],
  "three_lines_defence": {
    "first_line": {"title": "Continuous Control Monitoring", "description": "<desc>", "last_validated": "<date>"},
    "second_line": {"title": "Monthly RCSA (CSA)", "description": "<desc>", "last_validated": "<date>"},
    "third_line": {"title": "Internal Audit (Annual)", "description": "<desc>", "last_validated": "<date>"}
  },
  "priority_actions": [
    {"priority": "P1", "due_days": <integer>, "action": "<action description>", "owner": "<owner role>", "regulatory_ref": "<reg citation>"}
  ],
  "ai_rationale": "<paragraph explaining the scoring rationale, ~120-150 words>",
  "evidence_trace": ["<evidence item 1>", "<evidence item 2>", ...]
}

Base scores and content on the actual conversation. Include 5-6 detective controls, 3 priority actions, and 5-7 evidence trace items. Reference specific regulations cited in the conversation (GENIUS Act, FCA, FinCEN, etc.).

Return ONLY valid JSON. No markdown fences. No commentary."""


@app.post("/rcsa_output", response_model=RCSAOutputResponse)
async def rcsa_output(request: RCSAOutputRequest):
    """Generate structured RCSA output from interview conversation."""
    if not collection:
        raise HTTPException(503, "Knowledge base not initialised")

    if not request.messages:
        raise HTTPException(400, "No conversation provided")

    # Build conversation summary
    convo_text = "\n\n".join([
        f"{m.role.upper()}: {m.content}" for m in request.messages
    ])

    # Retrieve top relevant chunks across the whole conversation for grounding
    user_queries = " ".join([m.content for m in request.messages if m.role == "user"])
    context_block = ""
    if user_queries.strip():
        try:
            results = collection.query(
                query_texts=[user_queries[:1500]],
                n_results=6,
                include=["documents", "metadatas"],
            )
            if results["documents"] and results["documents"][0]:
                for doc, meta in zip(results["documents"][0], results["metadatas"][0]):
                    context_block += f"\n--- {meta['source']} ---\n{doc[:600]}\n"
        except Exception as e:
            log.error(f"RCSA output retrieval failed: {e}")

    try:
        completion = openai_client.chat.completions.create(
            model=CHAT_MODEL,
            messages=[
                {"role": "system", "content": RCSA_OUTPUT_PROMPT},
                {"role": "system", "content": f"RETRIEVED CONTEXT FROM KNOWLEDGE BASE:\n{context_block}"},
                {"role": "user", "content": f"BUSINESS UNIT: {request.business_unit}\n\nCONVERSATION TO ASSESS:\n\n{convo_text}\n\nGenerate the RCSA output as JSON."},
            ],
            max_tokens=2000,
            temperature=0.3,
            response_format={"type": "json_object"},
        )
        raw_json = completion.choices[0].message.content
        import json
        data = json.loads(raw_json)
    except Exception as e:
        log.error(f"RCSA output generation failed: {e}")
        raise HTTPException(500, f"RCSA output generation error: {str(e)}")

    return RCSAOutputResponse(
        business_unit=request.business_unit,
        inherent_risk_score=data.get("inherent_risk_score", 16),
        inherent_risk_label=data.get("inherent_risk_label", "HIGH"),
        residual_risk_score=data.get("residual_risk_score", 9),
        residual_risk_label=data.get("residual_risk_label", "MEDIUM-HIGH"),
        priority_action_count=data.get("priority_action_count", 3),
        heat_map_position=data.get("heat_map_position", {"impact": 4, "likelihood": 4}),
        detective_controls=data.get("detective_controls", []),
        three_lines_defence=data.get("three_lines_defence", {}),
        priority_actions=data.get("priority_actions", []),
        ai_rationale=data.get("ai_rationale", ""),
        evidence_trace=data.get("evidence_trace", []),
        generated_at=datetime.utcnow().isoformat(),
    )


# ── ENTERPRISE RISK DASHBOARD ────────────────────────────────────────────
DASHBOARD_PROMPT = """Generate a realistic enterprise risk dashboard for a large global retail bank's RCSA portfolio. Use the retrieved context to ground the data in real controls.

Return JSON ONLY with this EXACT structure:

{
  "business_units_in_scope": 12,
  "high_risk_bus": <integer 3-5>,
  "open_action_items": <integer 15-30>,
  "controls_rated_effective_pct": <integer 60-75>,
  "business_units": [
    {"name": "Retail Banking", "rating": "HIGH"},
    {"name": "Wealth Mgmt", "rating": "HIGH"},
    {"name": "Markets & Trading", "rating": "HIGH"},
    {"name": "Corp Lending", "rating": "MEDIUM"},
    {"name": "Digital Banking", "rating": "HIGH"},
    {"name": "Trade Finance", "rating": "MEDIUM"},
    {"name": "Compliance", "rating": "MEDIUM"},
    {"name": "Operations / IT", "rating": "LOW"},
    {"name": "HR", "rating": "LOW"},
    {"name": "Finance", "rating": "LOW"},
    {"name": "Treasury", "rating": "IN PROG"},
    {"name": "Legal", "rating": "IN PROG"}
  ],
  "top_risk_categories": [
    {"name": "Compliance / KYC / FCRA", "exposure_index": 91},
    {"name": "Fraud — Synthetic Identity", "exposure_index": 85},
    {"name": "Process & Data Execution", "exposure_index": 78},
    {"name": "Third Party / Vendor", "exposure_index": 68},
    {"name": "Technology / System", "exposure_index": 55},
    {"name": "AML / Financial Crime", "exposure_index": 42}
  ],
  "cycle_time_comparison": {
    "manual_days": 87,
    "ai_pilot_days": 31,
    "ai_current_days": 24,
    "reduction_pct": 72
  }
}

Return ONLY valid JSON. No markdown."""


@app.get("/dashboard", response_model=DashboardResponse)
async def dashboard():
    """Generate enterprise risk portfolio dashboard data."""
    if not collection:
        raise HTTPException(503, "Knowledge base not initialised")

    # Pull diverse context from the KB
    context_block = ""
    try:
        results = collection.query(
            query_texts=["operational risk control weakness across banking divisions"],
            n_results=8,
            include=["documents", "metadatas"],
        )
        if results["documents"] and results["documents"][0]:
            for doc, meta in zip(results["documents"][0], results["metadatas"][0]):
                context_block += f"\n--- {meta['source']} ---\n{doc[:400]}\n"
    except Exception as e:
        log.error(f"Dashboard retrieval failed: {e}")

    try:
        completion = openai_client.chat.completions.create(
            model=CHAT_MODEL,
            messages=[
                {"role": "system", "content": DASHBOARD_PROMPT},
                {"role": "system", "content": f"KNOWLEDGE BASE CONTEXT:\n{context_block}"},
                {"role": "user", "content": "Generate the dashboard JSON."},
            ],
            max_tokens=1500,
            temperature=0.4,
            response_format={"type": "json_object"},
        )
        import json
        data = json.loads(completion.choices[0].message.content)
    except Exception as e:
        log.error(f"Dashboard generation failed: {e}")
        raise HTTPException(500, f"Dashboard generation error: {str(e)}")

    return DashboardResponse(
        business_units_in_scope=data.get("business_units_in_scope", 12),
        high_risk_bus=data.get("high_risk_bus", 4),
        open_action_items=data.get("open_action_items", 23),
        controls_rated_effective_pct=data.get("controls_rated_effective_pct", 67),
        business_units=data.get("business_units", []),
        top_risk_categories=data.get("top_risk_categories", []),
        cycle_time_comparison=data.get("cycle_time_comparison", {}),
        generated_at=datetime.utcnow().isoformat(),
    )


# ── CONTROL GAP RECOMMENDATIONS ──────────────────────────────────────────
CONTROL_GAP_PROMPT = """You are a second-line risk reviewer analysing an RCSA interview transcript. Your job has two halves.

HALF 1 — UNMITIGATED RISKS
Identify every risk the control owner raised that the current control set does not adequately contain. A risk qualifies when the owner describes an exposure and either (a) names no control at all, (b) names a control they admit is manual, inconsistent, untested or breached, or (c) describes a workaround rather than a designed control. Only use risks actually voiced in the conversation — never invent one.

"Unmitigated" does NOT mean "no control exists". The usual case by far is that a relevant control DOES exist and simply does not reach this exposure — wrong trigger population, too infrequent, too high a threshold, executed manually, covering only part of the population. Such a risk is still unmitigated, and the existing control is still the right starting point for fixing it. Treat "nothing in the library relates to this at all" as the rare exception it is.

Before deciding, you MUST scan the AVAILABLE CONTROL INVENTORY for this specific risk and record that scan in "inventory_scan". Name the two or three closest controls by ID and say for each whether it bears on the risk. Any control your scan finds to bear on the risk — even partially — MUST then appear in "recommended_controls" as an EXTEND. Only write that nothing relates after you have actually looked and can name what you rejected.

For each risk, state what the risk team should DO about it. Every recommendation is one of exactly two actions:

  EXTEND — an existing control in the inventory already bears on this risk but does not reach far enough. Widening it is the recommendation. Put its real ID in "control_id" and its real inventory name in "name", say what it already does for this risk in "current_coverage", and say precisely what it fails to reach in "coverage_gap" (trigger population, scope, frequency, threshold, manual execution).
  ADD — nothing in the inventory bears on this risk, so a new control has to be built. Set "control_id" to null and leave "current_coverage" and "coverage_gap" as empty strings.

Work the AVAILABLE CONTROL INVENTORY before choosing. It lists every control's ID and what that control actually does — match on what a control does, never on its ID. EXTEND is the more common and more valuable answer: a control that bears on the risk only PARTIALLY is still the right thing to extend, and proposing a brand-new control that duplicates one already in the inventory is a failure of this task. Reach for ADD only when you have been through the inventory and nothing relates.

- NEVER put an ID in "control_id" unless it appears verbatim in the inventory, and never attach a name from the conversation to an inventory ID — a real ID under a wrong name is worse than no citation at all. The same applies to "control_name" in HALF 2.
- "why_unmitigated" may say that no control exists ONLY when every recommendation for that risk is an ADD.
- A risk may carry both: EXTEND the existing control AND ADD a new one alongside it.

HALF 2 — INADEQUATE CONTROLS
For every existing control the conversation shows to be unfit for the risk it is meant to mitigate, give the risk team both (a) how to test it and (b) what to change so it becomes adequate.

HALF 2 is independent of HALF 1 and is never optional. A control that you already cited in HALF 1 as EXISTING belongs here as well when the conversation shows it failing — listing it twice is correct and expected, because the two halves answer different questions. Never drop a control from HALF 2 to avoid repeating yourself. If the owner describes a control as manual, late, self-approved, unreviewed or breached, it belongs here even when it appears nowhere in the inventory (use the owner's own name for it as "control_id" and "control_name" in that case).

Return a JSON object with this EXACT structure — no markdown, no commentary:

{
  "unmitigated_risks": [
    {
      "risk_id": "UR-01",
      "risk_statement": "<the exposure, one sentence>",
      "risk_category": "<e.g. Fraud, AML, Process Execution, Third Party, Technology>",
      "severity": "<HIGH|MEDIUM|LOW>",
      "owner_statement": "<short quote or close paraphrase of what the control owner actually said>",
      "inventory_scan": "<the 2-3 closest inventory controls by ID, and for each whether it bears on this risk — e.g. 'RB-CTRL-02 covers synthetic identity but only post-decision on approved loans; RB-CTRL-04 covers KYC at onboarding, not origination screening'>",
      "why_unmitigated": "<why the current control set does not contain this risk — reference the controls named in inventory_scan rather than claiming nothing exists>",
      "regulatory_ref": "<citation if the conversation or context supports one, else empty string>",
      "recommended_controls": [
        {
          "action": "<EXTEND|ADD>",
          "control_id": "<real ID from the inventory when EXTEND, null when ADD>",
          "name": "<the control's real inventory name when EXTEND, the proposed name when ADD>",
          "control_type": "<Preventive|Detective|Corrective>",
          "objective": "<what the control should prevent or detect for this risk>",
          "design": "<for EXTEND: the change that widens it to cover this risk. For ADD: how the new control operates. 1-2 sentences>",
          "frequency": "<Continuous|Daily|Weekly|Monthly|Quarterly>",
          "owner": "<role that should own it>",
          "current_coverage": "<EXTEND only: what this control already does for this risk. ADD: empty string>",
          "coverage_gap": "<EXTEND only: what it fails to reach today. ADD: empty string>"
        }
      ]
    }
  ],
  "inadequate_controls": [
    {
      "control_id": "<real ID from the inventory, or the name the owner used>",
      "control_name": "<name>",
      "current_rating": "<Weak|Ineffective|Adequate but insufficient>",
      "why_inadequate": "<the specific design or operating deficiency>",
      "linked_risk_id": "<the UR-xx this control fails to mitigate, or empty string>",
      "test_procedure": {
        "objective": "<what the test proves>",
        "population": "<the population to test against>",
        "sample_size": "<e.g. 25 items, or 100% for low-volume populations>",
        "sampling_basis": "<Random|Judgmental|Risk-based|Full population — and why>",
        "test_attributes": ["<attribute to check on each sampled item>", "..."],
        "evidence_required": ["<artefact to request from the control owner>", "..."],
        "pass_criteria": "<what constitutes an effective result>",
        "fail_criteria": "<what constitutes a deficiency>"
      },
      "remediation": {
        "current_deficiency": "<one-line restatement of the gap>",
        "design_changes": ["<specific change that would raise this control to adequate>", "..."],
        "target_rating": "<Adequate|Strong>",
        "residual_risk_after": "<HIGH|MEDIUM-HIGH|MEDIUM|MEDIUM-LOW|LOW>",
        "owner": "<role accountable for the fix>",
        "due_days": <integer>
      }
    }
  ],
  "coverage_summary": "<paragraph of ~100-130 words on how much of the voiced risk the current control set actually covers, and the single biggest hole>",
  "evidence_trace": ["<what in the transcript or retrieved documents each conclusion rests on>", "..."]
}

Include 3-5 test_attributes and 2-4 design_changes per inadequate control, and 4-6 evidence_trace items. If the conversation genuinely surfaces no unmitigated risk, return empty arrays and say so in coverage_summary — do not manufacture findings.

Return ONLY valid JSON. No markdown fences. No commentary."""


def control_inventory() -> str:
    """One line per control: the ID plus what that control actually does.

    The description is what makes this usable. A bare list of IDs is a hallucination
    guard but not a matching aid — the model cannot map a risk onto an ID it has no
    description for, so it designs a new control instead of citing the real one, and
    any ID it does reach for gets an arbitrary name attached.
    """
    if not collection:
        return ""
    try:
        items = collection.get(
            where={"$and": [{"doc_type": "Control Narrative"}, {"chunk_index": 0}]},
            include=["documents", "metadatas"],
        )
    except Exception as e:
        log.error(f"Control inventory lookup failed: {e}")
        return ""

    lines = {}
    for doc, meta in zip(items["documents"], items["metadatas"]):
        control_id = meta.get("control_id")
        if control_id and control_id not in lines:
            summary = " ".join(doc.split())[:220]
            lines[control_id] = f"- {control_id}: {summary}"
    return "\n".join(lines[cid] for cid in sorted(lines))


@app.post("/control_recommendations", response_model=ControlGapResponse)
async def control_recommendations(request: ControlGapRequest):
    """Recommend controls for risks the interview left unmitigated, plus test procedures
    and remediation for controls shown to be unfit. Additive — /rcsa_output is untouched."""
    if not ENABLE_CONTROL_RECO:
        raise HTTPException(404, "Control recommendations are not enabled on this deployment")

    if not collection:
        raise HTTPException(503, "Knowledge base not initialised")

    if not request.messages:
        raise HTTPException(400, "No conversation provided")

    convo_text = "\n\n".join([
        f"{m.role.upper()}: {m.content}" for m in request.messages
    ])

    # Retrieve against the owner's own words — the risk language lives there
    user_queries = " ".join([m.content for m in request.messages if m.role == "user"])
    context_block = ""
    if user_queries.strip():
        try:
            results = collection.query(
                query_texts=[user_queries[:1500]],
                n_results=CONTROL_GAP_TOP_K,
                include=["documents", "metadatas"],
            )
            if results["documents"] and results["documents"][0]:
                for doc, meta in zip(results["documents"][0], results["metadatas"][0]):
                    context_block += f"\n--- {meta['source']} ({meta['control_id']}) ---\n{doc[:600]}\n"
        except Exception as e:
            log.error(f"Control gap retrieval failed: {e}")

    try:
        completion = openai_client.chat.completions.create(
            model=CHAT_MODEL,
            messages=[
                {"role": "system", "content": CONTROL_GAP_PROMPT},
                {"role": "system", "content": f"AVAILABLE CONTROL INVENTORY (the only IDs you may cite):\n{control_inventory()}"},
                {"role": "system", "content": f"RETRIEVED CONTEXT FROM KNOWLEDGE BASE:\n{context_block}"},
                {"role": "user", "content": f"BUSINESS UNIT: {request.business_unit}\n\nINTERVIEW TRANSCRIPT:\n\n{convo_text}\n\nIdentify the unmitigated risks and inadequate controls as JSON."},
            ],
            max_tokens=4000,
            temperature=0.3,
            response_format={"type": "json_object"},
        )
        import json
        data = json.loads(completion.choices[0].message.content)
    except Exception as e:
        log.error(f"Control gap generation failed: {e}")
        raise HTTPException(500, f"Control recommendation error: {str(e)}")

    unmitigated_risks = data.get("unmitigated_risks", []) or []
    inadequate_controls = data.get("inadequate_controls", []) or []
    log.info(f"Control gaps · {len(unmitigated_risks)} unmitigated risks · {len(inadequate_controls)} inadequate controls")

    return ControlGapResponse(
        business_unit=request.business_unit,
        unmitigated_risk_count=len(unmitigated_risks),
        inadequate_control_count=len(inadequate_controls),
        unmitigated_risks=unmitigated_risks,
        inadequate_controls=inadequate_controls,
        coverage_summary=data.get("coverage_summary", ""),
        evidence_trace=data.get("evidence_trace", []),
        generated_at=datetime.utcnow().isoformat(),
    )


# ── VOICE INPUT — AZURE SPEECH TOKEN ─────────────────────────────────────
@app.get("/speech_token")
async def speech_token():
    """Mint a short-lived Azure Speech token for the browser.

    The subscription key never leaves the server. The browser gets a 10-minute
    bearer token and streams microphone audio straight to the Azure Speech
    endpoint in the configured region — so audio goes to the tenant's own Speech
    resource rather than to a third-party dictation service.
    """
    if not VOICE_AVAILABLE:
        raise HTTPException(404, "Voice input is not enabled on this deployment")

    url = f"https://{AZURE_SPEECH_REGION}.api.cognitive.microsoft.com/sts/v1.0/issueToken"
    try:
        import httpx
        async with httpx.AsyncClient(timeout=10) as client:
            resp = await client.post(url, headers={
                "Ocp-Apim-Subscription-Key": AZURE_SPEECH_KEY,
                "Content-Length": "0",
            })
        resp.raise_for_status()
    except Exception as e:
        # Never echo the upstream body — it can carry key material back to the caller
        log.error(f"Speech token request failed: {type(e).__name__}: {e}")
        raise HTTPException(502, "Could not obtain a speech token from Azure Speech")

    return {
        "token": resp.text,
        "region": AZURE_SPEECH_REGION,
        "language": AZURE_SPEECH_LANGUAGE,
        "expires_in": SPEECH_TOKEN_TTL_SECONDS,
    }


if __name__ == "__main__":
    import uvicorn
    port = int(os.getenv("PORT", 8000))
    uvicorn.run(app, host="0.0.0.0", port=port)
